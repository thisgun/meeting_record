"""회의록 export — Word(.docx) / HTML 생성.

마크다운 본문을 파싱해서 docx의 헤더/리스트/굵게 처리.
지원하는 마크다운:
    # 제목
    ## / ### / #### 헤더
    - / * 리스트 (들여쓰기 지원)
    - [ ] / - [x] 체크박스
    **굵게**, *기울임*
    빈 줄 → 단락 구분
"""
from __future__ import annotations

import html
import re
from datetime import datetime
from pathlib import Path
from typing import Optional


def _format_timestamp(sec: float) -> str:
    sec = int(sec)
    h, rem = divmod(sec, 3600)
    m, s = divmod(rem, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def _format_duration(sec: float) -> str:
    if sec < 60:
        return f"{sec:.0f}초"
    if sec < 3600:
        m, s = divmod(int(sec), 60)
        return f"{m}분 {s}초"
    h, rem = divmod(int(sec), 3600)
    m = rem // 60
    return f"{h}시간 {m}분"


# ── DOCX ──────────────────────────────────────────────────────────

def to_docx(
    meeting: dict,
    utterances: list[dict],
    out_path: str | Path,
    *,
    include_transcript: bool = True,
    font_name: str = "맑은 고딕",
) -> Path:
    """회의 데이터를 .docx로 export.

    meeting: storage.get_meeting()["meeting"] 형식
    utterances: storage.get_meeting()["utterances"] 형식
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)

    # 1. 표지 정보
    title = meeting.get("title", "회의록")
    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.name = font_name

    # 메타 정보 표
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"
    meta = [
        ("작성 일시", meeting.get("created_at", "")),
        ("원본 파일", Path(meeting.get("source_file", "")).name),
        ("회의 길이", _format_duration(meeting.get("duration_sec", 0))),
        ("화자 수", f"{meeting.get('speaker_count', 0)}명"),
    ]
    for i, (k, v) in enumerate(meta):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = str(v)
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = font_name

    doc.add_paragraph()

    # 2. 요약 본문 (마크다운 파싱)
    md = meeting.get("summary_md", "")
    _render_markdown_to_docx(doc, md, font_name=font_name)

    # 3. 발화 전문 (선택)
    if include_transcript and utterances:
        doc.add_page_break()
        h = doc.add_heading("발화 전문", level=1)
        for run in h.runs:
            run.font.name = font_name

        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Light List Accent 1"
        hdr = tbl.rows[0].cells
        hdr[0].text = "시간"
        hdr[1].text = "화자"
        hdr[2].text = "내용"
        for cell in hdr:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = font_name
                    run.font.bold = True

        for u in utterances:
            row = tbl.add_row().cells
            row[0].text = _format_timestamp(u.get("start_sec", 0))
            row[1].text = u.get("speaker", "")
            row[2].text = u.get("text", "")
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = font_name

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def _render_markdown_to_docx(doc, md: str, *, font_name: str = "맑은 고딕") -> None:
    """간단한 마크다운 → docx 변환. 회의록에서 흔히 쓰는 요소만 지원."""
    from docx.shared import Pt

    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 빈 줄
        if not line.strip():
            i += 1
            continue

        # 헤더
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            h = doc.add_heading(m.group(2).strip(), level=min(level, 4))
            for run in h.runs:
                run.font.name = font_name
            i += 1
            continue

        # 리스트 (들여쓰기 지원, 체크박스)
        list_m = re.match(r"^(\s*)([\-\*])\s+(.*)", line)
        if list_m:
            indent = len(list_m.group(1)) // 2  # 2칸 = 1단계
            content = list_m.group(3).strip()

            # 체크박스
            cb_m = re.match(r"^\[([ x])\]\s+(.*)", content)
            if cb_m:
                checked = "☑" if cb_m.group(1) == "x" else "☐"
                content = f"{checked} {cb_m.group(2)}"

            style_name = "List Bullet" if indent == 0 else "List Bullet 2"
            try:
                p = doc.add_paragraph(style=style_name)
            except KeyError:
                p = doc.add_paragraph()
            _add_runs_with_bold(p, content, font_name=font_name)
            i += 1
            continue

        # 일반 단락 — 인접 비공백 줄 합치기
        para_lines = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip()
            if not nxt.strip():
                break
            if re.match(r"^(#{1,4})\s+", nxt) or re.match(r"^(\s*)([\-\*])\s+", nxt):
                break
            para_lines.append(nxt)
            j += 1
        text = " ".join(para_lines).strip()
        p = doc.add_paragraph()
        _add_runs_with_bold(p, text, font_name=font_name)
        i = j


def _add_runs_with_bold(paragraph, text: str, *, font_name: str = "맑은 고딕") -> None:
    """**굵게** 처리."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.name = font_name


# ── HTML ──────────────────────────────────────────────────────────

_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
body {{ font-family: '맑은 고딕', 'Malgun Gothic', sans-serif; max-width: 900px; margin: 2em auto; padding: 0 1em; line-height: 1.6; color: #222; }}
h1 {{ border-bottom: 3px solid #333; padding-bottom: 0.3em; }}
h2 {{ border-bottom: 1px solid #ccc; padding-bottom: 0.2em; margin-top: 1.5em; }}
h3 {{ margin-top: 1em; }}
.meta {{ background: #f4f4f4; padding: 1em; border-left: 4px solid #555; margin-bottom: 2em; }}
.meta dt {{ font-weight: bold; display: inline-block; width: 8em; }}
.meta dd {{ display: inline; margin: 0; }}
.meta div {{ margin-bottom: 0.4em; }}
.transcript {{ margin-top: 3em; border-top: 2px solid #333; padding-top: 1em; }}
.transcript table {{ border-collapse: collapse; width: 100%; }}
.transcript th, .transcript td {{ border: 1px solid #ddd; padding: 0.5em; text-align: left; vertical-align: top; }}
.transcript th {{ background: #eee; }}
.transcript .ts {{ white-space: nowrap; color: #666; font-family: 'Consolas', monospace; }}
.transcript .speaker {{ white-space: nowrap; font-weight: bold; }}
ul {{ margin: 0.5em 0; }}
li {{ margin: 0.2em 0; }}
@media print {{ body {{ max-width: none; margin: 0; }} }}
</style>
</head>
<body>
{body}
</body>
</html>"""


def to_html(
    meeting: dict,
    utterances: list[dict],
    out_path: str | Path,
    *,
    include_transcript: bool = True,
) -> Path:
    """회의 데이터를 HTML로 export (브라우저에서 Ctrl+P로 PDF 저장 가능)."""
    title = html.escape(meeting.get("title", "회의록"))
    body_parts: list[str] = [f"<h1>{title}</h1>"]

    body_parts.append('<div class="meta">')
    body_parts.append(f'<div><dt>작성 일시:</dt><dd>{html.escape(meeting.get("created_at",""))}</dd></div>')
    body_parts.append(f'<div><dt>원본 파일:</dt><dd>{html.escape(Path(meeting.get("source_file","")).name)}</dd></div>')
    body_parts.append(f'<div><dt>회의 길이:</dt><dd>{_format_duration(meeting.get("duration_sec",0))}</dd></div>')
    body_parts.append(f'<div><dt>화자 수:</dt><dd>{meeting.get("speaker_count",0)}명</dd></div>')
    body_parts.append("</div>")

    body_parts.append(_markdown_to_html(meeting.get("summary_md", "")))

    if include_transcript and utterances:
        body_parts.append('<div class="transcript">')
        body_parts.append("<h2>발화 전문</h2>")
        body_parts.append("<table><thead><tr><th>시간</th><th>화자</th><th>내용</th></tr></thead><tbody>")
        for u in utterances:
            body_parts.append("<tr>")
            body_parts.append(f'<td class="ts">{_format_timestamp(u.get("start_sec",0))}</td>')
            body_parts.append(f'<td class="speaker">{html.escape(u.get("speaker",""))}</td>')
            body_parts.append(f'<td>{html.escape(u.get("text",""))}</td>')
            body_parts.append("</tr>")
        body_parts.append("</tbody></table>")
        body_parts.append("</div>")

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(
        _HTML_TEMPLATE.format(title=title, body="\n".join(body_parts)),
        encoding="utf-8",
    )
    return out_path


def _markdown_to_html(md: str) -> str:
    """간단한 마크다운 → HTML."""
    out: list[str] = []
    lines = md.split("\n")
    in_list = False
    in_list_indent = 0

    def close_list():
        nonlocal in_list, in_list_indent
        while in_list_indent > 0:
            out.append("</ul>")
            in_list_indent -= 1
        if in_list:
            out.append("</ul>")
            in_list = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            close_list()
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            close_list()
            lvl = len(m.group(1))
            out.append(f"<h{lvl}>{_inline_md_to_html(m.group(2).strip())}</h{lvl}>")
            i += 1
            continue

        list_m = re.match(r"^(\s*)([\-\*])\s+(.*)", line)
        if list_m:
            indent = len(list_m.group(1)) // 2
            content = list_m.group(3).strip()
            cb_m = re.match(r"^\[([ x])\]\s+(.*)", content)
            if cb_m:
                checked = "☑" if cb_m.group(1) == "x" else "☐"
                content = f"{checked} {cb_m.group(2)}"
            if not in_list:
                out.append("<ul>")
                in_list = True
            while in_list_indent < indent:
                out.append("<ul>")
                in_list_indent += 1
            while in_list_indent > indent:
                out.append("</ul>")
                in_list_indent -= 1
            out.append(f"<li>{_inline_md_to_html(content)}</li>")
            i += 1
            continue

        # 일반 단락
        close_list()
        para = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip()
            if not nxt.strip() or re.match(r"^(#{1,4})\s+", nxt) or re.match(r"^(\s*)([\-\*])\s+", nxt):
                break
            para.append(nxt)
            j += 1
        out.append(f"<p>{_inline_md_to_html(' '.join(para).strip())}</p>")
        i = j

    close_list()
    return "\n".join(out)


def _inline_md_to_html(text: str) -> str:
    """**굵게**, *기울임* + HTML escape."""
    text = html.escape(text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<em>\1</em>", text)
    return text
